from docxtpl import DocxTemplate
import sys
import pymorphy2
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

morph = pymorphy2.MorphAnalyzer()
# Страница документа №1
EDUCATION_FORM = "Очная"

NAME_FILE = sys.argv[1]

INSTITUTE = sys.argv[2]
PRACTICE_KIND_ROD = sys.argv[3]
PRACTICE_KIND_ROD_SPLIT = PRACTICE_KIND_ROD.split()
PRACTICE_KIND_ROD_ARR = []
for i in PRACTICE_KIND_ROD_SPLIT:
    PRACTICE_KIND_ROD_PARSE = morph.parse(i)[0]
    PRACTICE_KIND_ROD_WORD = PRACTICE_KIND_ROD_PARSE.inflect({'gent'}).word.capitalize()
    PRACTICE_KIND_ROD_ARR.append(PRACTICE_KIND_ROD_WORD)
PRACTICE_KIND_ROD = ' '.join(PRACTICE_KIND_ROD_ARR)

PRACTICE_TYPE_ROD = sys.argv[4]
PRACTICE_TYPE_ROD_SPLIT = PRACTICE_TYPE_ROD.split()
PRACTICE_TYPE_ROD_ARR = []
for i in PRACTICE_TYPE_ROD_SPLIT:
    PRACTICE_TYPE_ROD_PARSE = morph.parse(i)[0]
    PRACTICE_TYPE_ROD_WORD = PRACTICE_TYPE_ROD_PARSE.inflect({'gent'}).word.capitalize()
    PRACTICE_TYPE_ROD_ARR.append(PRACTICE_TYPE_ROD_WORD)
PRACTICE_TYPE_ROD = ' '.join(PRACTICE_TYPE_ROD_ARR)

YEAR = sys.argv[5]
CODE_AND_PREPARATION_DIRECTION = sys.argv[6]

STUDENTS_COURSE = sys.argv[7]
STUDENTS_GROUP = sys.argv[8]
REPORT_YEAR = sys.argv[9]

# Страница документа №2
PRACTICE_DEADLINES = sys.argv[10]
ORDER_NUMBER_AND_DATE = sys.argv[11]
PRACTICE_KIND_IMEN = sys.argv[3]
PRACTICE_TYPE_IMEN = sys.argv[4]


STUDENT_FULLNAME_SUCCESS = sys.argv[12].split(',')
STUDENT_FULLNAME_FAILURE = sys.argv[13].split(',')
SUCCESS_STUDENT_COUNT = len(STUDENT_FULLNAME_SUCCESS)
FAILURE_STUDENT_COUNT = len(STUDENT_FULLNAME_FAILURE)
PRACTICE_PLACE = sys.argv[14]
USU_CHIEF_FULLNAME = sys.argv[15]
ORGANIZATION_CHIEF_FULLNAME = sys.argv[16]


doc = Document("../templates/report_document_template.docx")

for style in doc.styles:
    if style.type == 1:
        style.font.name = 'Times New Roman'

table = doc.add_table(rows=1, cols=7)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = ' № п/п'
hdr_cells[1].text = ' ФИО обучающегося'
hdr_cells[2].text = ' Место прохождения практики'
hdr_cells[3].text = ' Город прохождения практики'
hdr_cells[4].text = ' Вид договора (долгосрочный/ краткосрочный)'
hdr_cells[5].text = ' Оплачиваемая практика (да/нет)'
hdr_cells[6].text = ' ФИО руководителей практики от организации'

hdr_cells[0].paragraphs[0].runs[0].bold = True
hdr_cells[1].paragraphs[0].runs[0].bold = True
hdr_cells[2].paragraphs[0].runs[0].bold = True
hdr_cells[3].paragraphs[0].runs[0].bold = True
hdr_cells[4].paragraphs[0].runs[0].bold = True
hdr_cells[5].paragraphs[0].runs[0].bold = True
hdr_cells[6].paragraphs[0].runs[0].bold = True

i = 1
for student_fullname_success in STUDENT_FULLNAME_SUCCESS:
    row_cells = table.add_row().cells
    row_cells[0].text = " " + str(i)
    row_cells[1].text = " " + student_fullname_success
    row_cells[2].text = " " + PRACTICE_PLACE
    row_cells[3].text = " " + "г.Ханты-Мансийск"
    row_cells[4].text = " " + "-"
    row_cells[5].text = " " + "-"
    row_cells[6].text = " " + USU_CHIEF_FULLNAME
    i = i + 1

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

for paragraph in doc.paragraphs:
    if "ТУТ ДОЛЖНА БЫТЬ ТАБЛИЦА" in paragraph.text:
        parent = paragraph._element.getparent()
        index = parent.index(paragraph._element)
        parent.insert(index, table._tbl)
        parent.remove(paragraph._element)
        break

doc.save(NAME_FILE)


doc1 = Document(NAME_FILE)

for style in doc1.styles:
    if style.type == 1:
        style.font.name = 'Times New Roman'

table = doc1.add_table(rows=1, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = '№ п/п'
hdr_cells[1].text = 'ФИО обучающегося'
hdr_cells[2].text = 'Причина непрохождения практики'


hdr_cells[0].paragraphs[0].runs[0].bold = True
hdr_cells[1].paragraphs[0].runs[0].bold = True
hdr_cells[2].paragraphs[0].runs[0].bold = True

i = 1
for student_fullname_failure in STUDENT_FULLNAME_FAILURE:
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = student_fullname_failure
    row_cells[2].text = "Студент не выполнил минимальные требования"
    i = i + 1

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

for paragraph in doc1.paragraphs:
    if "ТУТ ДОЛЖНА БЫТЬ ТАБЛИЦАА" in paragraph.text:
        parent = paragraph._element.getparent()
        index = parent.index(paragraph._element)
        parent.insert(index, table._tbl)
        parent.remove(paragraph._element)
        break

doc1.save(NAME_FILE)

doc2 = DocxTemplate(NAME_FILE)


context = {"INSTITUTE": INSTITUTE, "PRACTICE_KIND_ROD": PRACTICE_KIND_ROD, "PRACTICE_TYPE_ROD": PRACTICE_TYPE_ROD,
           "YEAR": YEAR, "CODE_AND_PREPARATION_DIRECTION": CODE_AND_PREPARATION_DIRECTION, "EDUCATION_FORM": EDUCATION_FORM,
           "STUDENTS_COURSE": STUDENTS_COURSE, "STUDENTS_GROUP": STUDENTS_GROUP, "REPORT_YEAR": REPORT_YEAR,
           "PRACTICE_DEADLINES": PRACTICE_DEADLINES, "ORDER_NUMBER_AND_DATE": ORDER_NUMBER_AND_DATE, "PRACTICE_KIND_IMEN": PRACTICE_KIND_IMEN,
           "PRACTICE_TYPE_IMEN": PRACTICE_TYPE_IMEN, "SUCCESS_STUDENT_COUNT": SUCCESS_STUDENT_COUNT, "FAILURE_STUDENT_COUNT": FAILURE_STUDENT_COUNT,
           "USU_CHIEF_FULLNAME": USU_CHIEF_FULLNAME, "ORGANIZATION_CHIEF_FULLNAME": ORGANIZATION_CHIEF_FULLNAME}
doc2.render(context)
doc2.save(NAME_FILE)





