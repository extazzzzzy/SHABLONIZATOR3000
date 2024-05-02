from docxtpl import DocxTemplate
import sys
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
import pymorphy2

morph = pymorphy2.MorphAnalyzer()

NAME_FILE = sys.argv[1]
PRACTICE_KIND_IMEN = sys.argv[2]
PRACTICE_KIND_DAT = sys.argv[3]
PRACTICE_KIND_VIN = sys.argv[4]
STUDENT_COURSE = sys.argv[5]
STUDENT_GROUP = sys.argv[6]
STUDENT_FULLNAME_IMEN = sys.argv[7]
STUDENT_FULLNAME_ROD = sys.argv[8]
STUDENT_FULLNAME_DAT = sys.argv[9]
INSTITUTE = sys.argv[10]
PREPARATION_DIRECTION = sys.argv[11]
USU_CHIEF_FULLNAME = sys.argv[12]
USU_CHIEF_POSITION = sys.argv[13]
ORGANIZATION_CHIEF_FULLNAME = sys.argv[14]
ORGANIZATION_CHIEF_POSITION = sys.argv[15]

PRACTICE_PLACE = sys.argv[16]
PRACTICE_PLACE_PRED = PRACTICE_PLACE
IS_PRACTICE_PLACE_PRED = False
PRACTICE_PLACE_PRED_SPLIT = PRACTICE_PLACE_PRED.split()
PRACTICE_PLACE_PRED_ARR = []
for i in PRACTICE_PLACE_PRED_SPLIT:
    if (IS_PRACTICE_PLACE_PRED == False):
        PRACTICE_PLACE_PRED_PARSE = morph.parse(i)[0]
        PRACTICE_PLACE_PRED_WORD = PRACTICE_PLACE_PRED_PARSE.inflect({'loct'}).word.capitalize()
        IS_PRACTICE_PLACE_PRED = True
    else:
        PRACTICE_PLACE_PRED_PARSE = morph.parse(i)[0]
        PRACTICE_PLACE_PRED_WORD = PRACTICE_PLACE_PRED_PARSE.inflect({'loct'}).word
    PRACTICE_PLACE_PRED_ARR.append(PRACTICE_PLACE_PRED_WORD)
PRACTICE_PLACE_PRED = ' '.join(PRACTICE_PLACE_PRED_ARR)

PRACTICE_PLACE_ADDRESS = sys.argv[17]
WORK_YEAR = sys.argv[18]
PRACTICE_DEADLINES = sys.argv[19]

taskNames = sys.argv[20].split(',')
taskDescriptions = sys.argv[21].split(',')



FULL_PATH_TO_TEMP_DOCUMENT = NAME_FILE + ".docx"
FULL_PATH_TO_MAIN_DOCUMENT = "../documents/" + FULL_PATH_TO_TEMP_DOCUMENT

doc = Document("../templates/diary_template.docx")

for style in doc.styles:
    if style.type == 1:
        style.font.name = 'Times New Roman'

table = doc.add_table(rows=2, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Выполненные виды работ в рамках задач (мероприятий), входящих в задание студента на практику'
hdr_cells[2].text = 'Подпись руководителя практики от организации'
hdr_cells[0].paragraphs[0].runs[0].bold = True
hdr_cells[2].paragraphs[0].runs[0].bold = True

table.cell(0, 0).merge(table.cell(0, 1))
table.cell(0, 2).merge(table.cell(1, 2))

hdr_cells = table.rows[1].cells
hdr_cells[0].text = 'Дата'
hdr_cells[1].text = 'Наименование работы'
hdr_cells[0].paragraphs[0].runs[0].bold = True
hdr_cells[1].paragraphs[0].runs[0].bold = True

for name, description in zip(taskNames, taskDescriptions):
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = description.capitalize()

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

for paragraph in doc.paragraphs:
    if "ТУТ ДОЛЖНА БЫТЬ ТАБЛИЦА" in paragraph.text:
        parent = paragraph._element.getparent()
        index = parent.index(paragraph._element)
        parent.insert(index, table._tbl)
        parent.remove(paragraph._element)
        break

doc.save(FULL_PATH_TO_TEMP_DOCUMENT)

doc1 = DocxTemplate(FULL_PATH_TO_TEMP_DOCUMENT)
context = {"PRACTICE_KIND_IMEN": PRACTICE_KIND_IMEN, "PRACTICE_KIND_DAT": PRACTICE_KIND_DAT, "PRACTICE_KIND_VIN":
    PRACTICE_KIND_VIN, "STUDENT_COURSE": STUDENT_COURSE, "STUDENT_GROUP": STUDENT_GROUP, "STUDENT_FULLNAME_IMEN":
    STUDENT_FULLNAME_IMEN, "STUDENT_FULLNAME_ROD": STUDENT_FULLNAME_ROD,
           "STUDENT_FULLNAME_DAT": STUDENT_FULLNAME_DAT, "INSTITUTE": INSTITUTE, "PREPARATION_DIRECTION": PREPARATION_DIRECTION,
           "USU_CHIEF_FULLNAME": USU_CHIEF_FULLNAME, "USU_CHIEF_POSITION": USU_CHIEF_POSITION, "TASKS": taskNames, "ORGANIZATION_CHIEF_FULLNAME": ORGANIZATION_CHIEF_FULLNAME,
           "ORGANIZATION_CHIEF_POSITION": ORGANIZATION_CHIEF_POSITION, "PRACTICE_PLACE": PRACTICE_PLACE, "PRACTICE_PLACE_PRED": PRACTICE_PLACE_PRED, "PRACTICE_PLACE_ADDRESS": PRACTICE_PLACE_ADDRESS, "WORK_YEAR": WORK_YEAR,
           "PRACTICE_DEADLINES": PRACTICE_DEADLINES,


           "STUDENT_QUALITIES": "{{STUDENT_QUALITIES}}", "PROBLEM_SOLVING_SPEED": "{{PROBLEM_SOLVING_SPEED}}",
           "WORK_AMOUNT": "{{WORK_AMOUNT}}", "REMARKS": "{{REMARKS}}",
           "STUDENT_ASSESSMENT": "{{STUDENT_ASSESSMENT}}"}
doc1.render(context)
doc1.save(FULL_PATH_TO_MAIN_DOCUMENT)
os.remove(FULL_PATH_TO_TEMP_DOCUMENT)