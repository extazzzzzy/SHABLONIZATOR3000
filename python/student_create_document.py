from docxtpl import DocxTemplate
import sys
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

NAME_FILE = sys.argv[1]
PRACTICE_KIND_IMEN = sys.argv[2]
PRACTICE_KIND_DAT = sys.argv[3]
PRACTICE_KIND_VIN = sys.argv[4]
STUDENT_COURSE = sys.argv[5]
STUDENT_GROUP = sys.argv[6]
STUDENT_FULLNAME_IMEN = sys.argv[7]
STUDENT_FULLNAME_ROD = sys.argv[8]
STUDENT_FULLNAME_DAT = sys.argv[9]
INSTITUTE = sys.argv[10]
PREPARATION_DIRECTION = sys.argv[11]
USU_CHIEF_FULLNAME = sys.argv[12]
USU_CHIEF_POSITION = sys.argv[13]

FULL_PATH_TO_TABLES_CSV = NAME_FILE + ".csv"
FULL_PATH_TO_TEMP_DOCUMENT = NAME_FILE + ".docx"
FULL_PATH_TO_MAIN_DOCUMENT = "../documents/" + FULL_PATH_TO_TEMP_DOCUMENT
# Вставка таблицы в документ
with open(FULL_PATH_TO_TABLES_CSV, newline='', encoding='utf-8') as csvfile:
    reader = csv.reader(csvfile)
    next(reader)

    doc = Document("../templates/diary_template.docx")

    for style in doc.styles:
        if style.type == 1:
            style.font.name = 'Times New Roman'

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'


    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Выполненные виды работ в рамках задач (мероприятий), входящих в задание студента на практику'
    hdr_cells[1].text = ''
    hdr_cells[2].text = 'Подпись руководителя практики от организации'
    hdr_cells[0].merge(hdr_cells[1])
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    hdr_cells[2].paragraphs[0].runs[0].bold = True

    new_row = table.add_row().cells
    new_row[0].text = 'Дата'
    new_row[1].text = 'Наименование работы'
    new_row[2].text = ''

    new_row[0].paragraphs[0].runs[0].bold = True
    new_row[1].paragraphs[0].runs[0].bold = True
    new_row[2].paragraphs[0].runs[0].bold = True

    for row in reader:
        row_cells = table.add_row().cells
        row_cells[0].text = date_without_time = row[2][:-9]
        row_cells[1].text = row[1]
        row_cells[2].text = ''


for paragraph in doc.paragraphs:
        if "ТУТ ДОЛЖНА БЫТЬ ТАБЛИЦА" in paragraph.text:
            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)
            parent.insert(index, table._tbl)

            parent.remove(paragraph._element)
            break

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save(FULL_PATH_TO_TEMP_DOCUMENT)

doc1 = DocxTemplate(FULL_PATH_TO_TEMP_DOCUMENT)
context = {"PRACTICE_KIND_IMEN": PRACTICE_KIND_IMEN, "PRACTICE_KIND_DAT": PRACTICE_KIND_DAT, "PRACTICE_KIND_VIN":
    PRACTICE_KIND_VIN, "STUDENT_COURSE": STUDENT_COURSE, "STUDENT_GROUP": STUDENT_GROUP, "STUDENT_FULLNAME_IMEN":
    STUDENT_FULLNAME_IMEN, "STUDENT_FULLNAME_ROD": STUDENT_FULLNAME_ROD,
           "STUDENT_FULLNAME_DAT": STUDENT_FULLNAME_DAT, "INSTITUTE": INSTITUTE, "PREPARATION_DIRECTION": PREPARATION_DIRECTION,
           "USU_CHIEF_FULLNAME": USU_CHIEF_FULLNAME, "USU_CHIEF_POSITION": USU_CHIEF_POSITION,

           "PRACTICE_PLACE": "{{PRACTICE_PLACE}}", "ORGANIZATION_CHIEF_FULLNAME": "{{ORGANIZATION_CHIEF_FULLNAME}}", "ORGANIZATION_CHIEF_POSITION": "{{ORGANIZATION_CHIEF_POSITION}}",
           "WORK_YEAR": "{{WORK_YEAR}}", "PRACTICE_DEADLINES": "{{PRACTICE_DEADLINES}}", "PRACTICE_PLACE_ADDRESS": "{{PRACTICE_PLACE_ADDRESS}}",
           "STUDENT_QUALITIES": "{{STUDENT_QUALITIES}}", "PROBLEM_SOLVING_SPEED": "{{PROBLEM_SOLVING_SPEED}}", "WORK_AMOUNT": "{{WORK_AMOUNT}}", "REMARKS": "{{REMARKS}}",
           "STUDENT_ASSESSMENT": "{{STUDENT_ASSESSMENT}}"}

doc1.render(context)
doc1.save(FULL_PATH_TO_MAIN_DOCUMENT)
os.remove(FULL_PATH_TO_TEMP_DOCUMENT)
os.remove(FULL_PATH_TO_TABLES_CSV)