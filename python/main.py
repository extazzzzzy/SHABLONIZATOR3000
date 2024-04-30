from docxtpl import DocxTemplate
import sys
import pymorphy2
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

morph = pymorphy2.MorphAnalyzer()
# ПОЛУЧЕНИЕ ВСЕХ ДАННЫХ ИЗ PHP
PRACTICE_PLACE = sys.argv[1]

# перевод места практики в предложный падеж
PRACTICE_PLACE_PRED = PRACTICE_PLACE
IS_PRACTICE_PLACE_PRED = False
PRACTICE_PLACE_PRED_SPLIT = PRACTICE_PLACE_PRED.split()
PRACTICE_PLACE_PRED_ARR = []
for i in PRACTICE_PLACE_PRED_SPLIT:
    if (IS_PRACTICE_PLACE_PRED == False):
        PRACTICE_PLACE_PRED_PARSE = morph.parse(i)[0]
        PRACTICE_PLACE_PRED_WORD = PRACTICE_PLACE_PRED_PARSE.inflect({'loct'}).word.capitalize()
        IS_PRACTICE_PLACE_PRED = True
    else:
        PRACTICE_PLACE_PRED_PARSE = morph.parse(i)[0]
        PRACTICE_PLACE_PRED_WORD = PRACTICE_PLACE_PRED_PARSE.inflect({'loct'}).word
    PRACTICE_PLACE_PRED_ARR.append(PRACTICE_PLACE_PRED_WORD)
PRACTICE_PLACE_PRED = ' '.join(PRACTICE_PLACE_PRED_ARR)

PRACTICE_PLACE_ADDRESS = sys.argv[2]
STUDENT_COURSE = sys.argv[3]
STUDENT_GROUP = sys.argv[4]

# перевод имени студента в родительный падеж
STUDENT_FULLNAME_ROD = sys.argv[5]
STUDENT_FULLNAME_ROD_SPLIT = STUDENT_FULLNAME_ROD.split()
STUDENT_FULLNAME_ROD_ARR = []
for i in STUDENT_FULLNAME_ROD_SPLIT:
    STUDENT_FULLNAME_ROD_PARSE = morph.parse(i)[0]
    STUDENT_FULLNAME_ROD_WORD = STUDENT_FULLNAME_ROD_PARSE.inflect({'gent'}).word.capitalize()
    STUDENT_FULLNAME_ROD_ARR.append(STUDENT_FULLNAME_ROD_WORD)
STUDENT_FULLNAME_ROD = ' '.join(STUDENT_FULLNAME_ROD_ARR)

STUDENT_FULLNAME_IMEN = sys.argv[5]

# перевод имени студента в дательный падеж
STUDENT_FULLNAME_DAT = STUDENT_FULLNAME_IMEN
STUDENT_FULLNAME_DAT_SPLIT = STUDENT_FULLNAME_DAT.split()
STUDENT_FULLNAME_DAT_ARR = []
for i in STUDENT_FULLNAME_DAT_SPLIT:
    STUDENT_FULLNAME_DAT_PARSE = morph.parse(i)[0]
    STUDENT_FULLNAME_DAT_WORD = STUDENT_FULLNAME_DAT_PARSE.inflect({'datv'}).word.capitalize()
    STUDENT_FULLNAME_DAT_ARR.append(STUDENT_FULLNAME_DAT_WORD)
STUDENT_FULLNAME_DAT = ' '.join(STUDENT_FULLNAME_DAT_ARR)

PRACTICE_KIND_IMEN = sys.argv[6]

# перевод вида практики в дательный падеж
PRACTICE_KIND_DAT = PRACTICE_KIND_IMEN
PRACTICE_KIND_DAT_SPLIT = PRACTICE_KIND_DAT.split()
PRACTICE_KIND_DAT_ARR = []
for i in PRACTICE_KIND_DAT_SPLIT:
    PRACTICE_KIND_DAT_PARSE = morph.parse(i)[0]
    PRACTICE_KIND_DAT_WORD = PRACTICE_KIND_DAT_PARSE.inflect({'datv'}).word.upper()
    PRACTICE_KIND_DAT_ARR.append(PRACTICE_KIND_DAT_WORD)
PRACTICE_KIND_DAT = ' '.join(PRACTICE_KIND_DAT_ARR)

# перевод вида практики в винительный падеж
PRACTICE_KIND_VIN = PRACTICE_KIND_IMEN
PRACTICE_KIND_VIN_SPLIT = PRACTICE_KIND_VIN.split()
PRACTICE_KIND_VIN_ARR = []
for i in PRACTICE_KIND_VIN_SPLIT:
    PRACTICE_KIND_VIN_PARSE = morph.parse(i)[0]
    PRACTICE_KIND_VIN_WORD = PRACTICE_KIND_VIN_PARSE.inflect({'accs'}).word.capitalize()
    PRACTICE_KIND_VIN_ARR.append(PRACTICE_KIND_VIN_WORD)
PRACTICE_KIND_VIN = ' '.join(PRACTICE_KIND_VIN_ARR)

ORGANIZATION_CHIEF_FULLNAME = sys.argv[7]
ORGANIZATION_CHIEF_POSITION = sys.argv[8]
USU_CHIEF_FULLNAME = sys.argv[9]
USU_CHIEF_POSITION = sys.argv[10]

INSTITUTE = sys.argv[11]
PRACTICE_DEADLINES = sys.argv[12]
WORK_YEAR = sys.argv[13]
PREPARATION_DIRECTION = sys.argv[14]
STUDENT_QUALITIES = sys.argv[15]
PROBLEM_SOLVING_SPEED = sys.argv[16]
WORK_AMOUNT = sys.argv[17]
REMARKS = sys.argv[18]
STUDENT_ASSESSMENT = sys.argv[19]

# внедрение .csv таблицы в документ
with open('tables.csv', newline='', encoding='utf-8') as csvfile:
    reader = csv.reader(csvfile)
    next(reader)

    doc = Document("../templates/diary_template.docx")

    for style in doc.styles:
        if style.type == 1:
            style.font.name = 'Times New Roman'

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'


    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Выполненные виды работ в рамках задач (мероприятий), входящих в задание студента на практику'
    hdr_cells[1].text = ''
    hdr_cells[2].text = 'Подпись руководителя практики от организации'
    hdr_cells[0].merge(hdr_cells[1])
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    hdr_cells[2].paragraphs[0].runs[0].bold = True

    new_row = table.add_row().cells
    new_row[0].text = 'Дата'
    new_row[1].text = 'Наименование работы'
    new_row[2].text = ''

    new_row[0].paragraphs[0].runs[0].bold = True
    new_row[1].paragraphs[0].runs[0].bold = True
    new_row[2].paragraphs[0].runs[0].bold = True

    for row in reader:
        row_cells = table.add_row().cells
        row_cells[0].text = date_without_time = row[2][:-9]
        row_cells[1].text = row[1]
        row_cells[2].text = ''


for paragraph in doc.paragraphs:
        if "ТУТ ДОЛЖНА БЫТЬ ТАБЛИЦА" in paragraph.text:
            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)
            parent.insert(index, table._tbl)

            parent.remove(paragraph._element)
            break

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save("temp.docx")
doc1 = DocxTemplate("temp.docx")
context = {"PRACTICE_PLACE": PRACTICE_PLACE, "PRACTICE_PLACE_PRED": PRACTICE_PLACE_PRED, "PRACTICE_PLACE_ADDRESS":
    PRACTICE_PLACE_ADDRESS, "STUDENT_COURSE": STUDENT_COURSE, "STUDENT_GROUP": STUDENT_GROUP, "STUDENT_FULLNAME_ROD":
    STUDENT_FULLNAME_ROD, "STUDENT_FULLNAME_IMEN": STUDENT_FULLNAME_IMEN, "STUDENT_FULLNAME_DAT": STUDENT_FULLNAME_DAT, "PRACTICE_KIND_IMEN": PRACTICE_KIND_IMEN, "PRACTICE_KIND_DAT": PRACTICE_KIND_DAT, "PRACTICE_KIND_VIN": PRACTICE_KIND_VIN, "ORGANIZATION_CHIEF_FULLNAME": ORGANIZATION_CHIEF_FULLNAME, "ORGANIZATION_CHIEF_POSITION": ORGANIZATION_CHIEF_POSITION, "USU_CHIEF_FULLNAME": USU_CHIEF_FULLNAME, "USU_CHIEF_POSITION": USU_CHIEF_POSITION, "INSTITUTE": INSTITUTE, "PRACTICE_DEADLINES": PRACTICE_DEADLINES, "WORK_YEAR": WORK_YEAR, "PREPARATION_DIRECTION": PREPARATION_DIRECTION, "STUDENT_QUALITIES": STUDENT_QUALITIES, "PROBLEM_SOLVING_SPEED": PROBLEM_SOLVING_SPEED, "WORK_AMOUNT": WORK_AMOUNT, "REMARKS": REMARKS, "STUDENT_ASSESSMENT": STUDENT_ASSESSMENT}
doc1.render(context)
doc1.save("../documents/ready_document.docx")
os.remove("temp.docx")